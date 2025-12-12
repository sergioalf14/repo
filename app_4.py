# app_4.py (full persistence, single-click navigation, steps 1-9)
import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
import base64
import requests
import tempfile
import traceback

# ------------------------------------------------
# CONFIG
# ------------------------------------------------
ALIGNMENT_FILE = "strategic_alignment.xlsx"

# Candidate persistent dir on Streamlit Cloud
DATA_DIR_CANDIDATE = "/mount/data/workplan_data"

# GitHub defaults (you provided repo)
DEFAULT_GITHUB_REPO = "sergioalf14/repo"

USE_GITHUB = True   # Keep GitHub syncing enabled (Option A)

# ------------------------------------------------
# SAFE DATA DIRECTORY INITIALIZATION
# ------------------------------------------------
try:
    os.makedirs(DATA_DIR_CANDIDATE, exist_ok=True)
    LOCAL_DATA_DIR = DATA_DIR_CANDIDATE
except Exception:
    LOCAL_DATA_DIR = tempfile.gettempdir()

# Ensure final MASTER_LOG path is defined after LOCAL_DATA_DIR is final
MASTER_LOG = os.path.join(LOCAL_DATA_DIR, "master_log.xlsx")
ANNEX_DIR = os.path.join(LOCAL_DATA_DIR, "annexes")
os.makedirs(ANNEX_DIR, exist_ok=True)

# ------------------------------------------------
# Load Streamlit secrets (allow override)
# ------------------------------------------------
GITHUB_TOKEN = st.secrets.get("GITHUB_TOKEN", None) if hasattr(st, "secrets") else None
GITHUB_REPO = st.secrets.get("GITHUB_REPO", DEFAULT_GITHUB_REPO) if hasattr(st, "secrets") else DEFAULT_GITHUB_REPO
GITHUB_BRANCH = st.secrets.get("GITHUB_BRANCH", "main") if hasattr(st, "secrets") else "main"

# If user disabled GitHub, allow override
if not USE_GITHUB:
    GITHUB_TOKEN = None

# ------------------------------------------------
# Helper: Push a file to GitHub (safe)
# ------------------------------------------------
def push_file_to_github(local_path, github_path):
    """
    Uploads (creates or updates) a file to the GitHub repo path.
    Returns a tuple: (success: bool, message: str)
    """
    if not USE_GITHUB:
        return (False, "GitHub upload disabled by configuration.")
    if not GITHUB_TOKEN:
        return (False, "GITHUB_TOKEN missing in Streamlit secrets.")
    if not GITHUB_REPO:
        return (False, "GITHUB_REPO not configured.")

    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{github_path}"
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
        "Accept": "application/vnd.github+json",
    }

    try:
        with open(local_path, "rb") as f:
            content_b64 = base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        return (False, f"Failed to read local file for upload: {e}")

    # Check if file already exists to get sha
    try:
        r = requests.get(url, headers=headers, timeout=30)
    except Exception as e:
        return (False, f"GitHub API GET failed: {e}")

    sha = None
    if r.status_code == 200:
        try:
            resp = r.json()
            sha = resp.get("sha")
        except Exception:
            sha = None
    elif r.status_code not in (404,):
        # Unexpected error
        return (False, f"GitHub API GET error: {r.status_code} - {r.text}")

    commit_msg = f"Update {github_path}" if sha else f"Add {github_path}"
    payload = {
        "message": commit_msg,
        "content": content_b64,
        "branch": GITHUB_BRANCH
    }
    if sha:
        payload["sha"] = sha

    try:
        r2 = requests.put(url, json=payload, headers=headers, timeout=60)
    except Exception as e:
        return (False, f"GitHub API PUT failed: {e}")

    if r2.status_code in (200, 201):
        return (True, f"Pushed to GitHub: {github_path}")
    else:
        # return API error text to help debugging
        return (False, f"GitHub upload failed: {r2.status_code} - {r2.text}")

# ------------------------------------------------
# FILE EXISTENCE CHECK
# ------------------------------------------------
if not os.path.exists(ALIGNMENT_FILE):
    st.error(" Missing file: strategic_alignment.xlsx — please add it to your repository.")
    st.stop()

# ------------------------------------------------
# SESSION STATE INITIALIZATION
# ------------------------------------------------
if "step" not in st.session_state:
    st.session_state.step = 1
if "submission" not in st.session_state:
    st.session_state.submission = {}
if "last_file" not in st.session_state:
    st.session_state.last_file = None
if "annex_saved_list" not in st.session_state:
    st.session_state.annex_saved_list = []  # list of tuples (filename, saved_path)

# ----------------------------
# NAVIGATION FUNCTIONS
# ----------------------------
def next_step():
    st.session_state.step += 1

def prev_step():
    if st.session_state.step > 1:
        st.session_state.step -= 1


# ------------------------------------------------
# Save to master Excel (robust and append-only)
# ------------------------------------------------
def save_to_master_excel(row_dict):
    """
    Safely append a row to MASTER_LOG (xlsx). Falls back to tempdir if needed.
    Also attempts to push to GitHub (master_log.xlsx at repo root).
    Returns (success: bool, message: str)
    """

    local_master = MASTER_LOG

    # ------------------------------------------------
    # 1. Always build the full DataFrame BEFORE writing
    # ------------------------------------------------
    df_new = pd.DataFrame([row_dict])  # new row must be list of dict

    if os.path.exists(local_master):
        try:
            df_old = pd.read_excel(local_master)
        except Exception:
            df_old = pd.DataFrame()
    else:
        df_old = pd.DataFrame()

    # Append new entry
    df_final = pd.concat([df_old, df_new], ignore_index=True)

    # ------------------------------------------------
    # 2. Try writing to the actual location
    # ------------------------------------------------
    try:
        df_final.to_excel(local_master, index=False)

    except PermissionError:
        # ------------------------------------------------
        # 3. Permission error → save to temp directory fallback
        # ------------------------------------------------
        fallback = os.path.join(tempfile.gettempdir(), "master_log.xlsx")
        try:
            df_final.to_excel(fallback, index=False)
            local_master = fallback
        except Exception as e:
            return (False, f"Failed to write master log even to fallback: {e}")

    except Exception as e:
        return (False, f"Failed to write master log: {e}")

    # ------------------------------------------------
    # 4. Push to GitHub (if enabled)
    # ------------------------------------------------
    success, msg = push_file_to_github(local_master, "master_log.xlsx")

    if success:
        return (True, "")
    else:
        # Local write succeeded, GitHub failed → still OK
        return (True, f"Master log written locally at {local_master}. GitHub push: {msg}")

# ------------------------------------------------
# Export Word (generate + push + return filepath)
# ------------------------------------------------
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Pt(24)
    return p

def add_bullet_list(doc, items):
    for x in items:
        if x:
            p = doc.add_paragraph(x, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)

def add_table_from_dict(doc, d):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in d.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value if value else ""
    doc.add_paragraph("")  # spacing

def export_word(data):
    """
    Creates a fully formatted WTO-style Word document, saves it to disk,
    attempts to push to GitHub, and returns (filepath, filename, push_result_msg).
    """
    try:
        doc = Document()

        # -----------------
        # Build document (same as your current logic)
        # -----------------
        add_heading(doc, "Divisional Workplan Summary", level=1)
        cover = data.get("Cover", {})
        for k in ["Division", "Director", "Date", "Version", "FTEs", "Financial Resources", "Director Signature"]:
            add_paragraph(doc, f"{k}: {cover.get(k, '')}")

        doc.add_page_break()

        # 1. Strategic Goals
        add_heading(doc, "1. Strategic Goals")
        add_bullet_list(doc, data.get("Selected Goals", []))

        # 2. Aggregate Objectives
        add_heading(doc, "2. Aggregate Divisional Objectives")
        for goal, objs in data.get("Aggregate Objectives", {}).items():
            add_heading(doc, f"{goal}", level=2)
            add_bullet_list(doc, objs)

        # 3. Activities & Expected Results
        add_heading(doc, "3. Activities and Expected Results")
        for (goal, agg), content in data.get("Activities", {}).items():
            add_heading(doc, f"{goal} — {agg}", level=2)

            if content.get("activities"):
                add_paragraph(doc, "Activities:", bold=True)
                add_bullet_list(doc, content["activities"])

            if content.get("results"):
                add_paragraph(doc, "Expected Results:", bold=True)
                add_bullet_list(doc, content["results"])

        # 4. Goal Metrics
        add_heading(doc, "4. Metrics for Strategic Goals")
        for goal, m in data.get("Goal Metrics", {}).items():
            add_heading(doc, goal, level=2)
            add_table_from_dict(doc, {
                "FTEs": m.get("FTEs", ""),
                "Financial Resources": m.get("Financial Resources", ""),
                "KPIs": m.get("KPIs", ""),
                "Other Metrics": m.get("Other Metrics", "")
            })

        # 5. Objective/Result Metrics
        add_heading(doc, "5. Metrics for Objectives and Results")
        for (goal, agg, tag), m in data.get("Objective/Result Metrics", {}).items():
            if tag == "AGGREGATE":
                title = f"{goal} — {agg}"
            else:
                title = f"Expected Result: {tag[4:]}"
            add_heading(doc, title, level=2)

            add_table_from_dict(doc, {
                "FTEs": m.get("FTEs", ""),
                "Financial Resources": m.get("Financial Resources", ""),
                "KPIs": m.get("KPIs", ""),
                "Other Metrics": m.get("Other Metrics", "")
            })

        # 6. Additional Information
        add_heading(doc, "6. Additional Information")
        for k, v in data.get("Additional", {}).items():
            add_paragraph(doc, k + ":", bold=True)
            add_paragraph(doc, v if v else "—", indent=True)

        # 7. Annexes
        add_heading(doc, "7. Annexes")
        annex_items = data.get("Annexes_Saved", [])
        safe_names = []
        for item in annex_items:
            if isinstance(item, dict) and "original_name" in item:
                safe_names.append(item["original_name"])
            elif isinstance(item, str):
                safe_names.append(os.path.basename(item))
            else:
                safe_names.append(str(item))

        if safe_names:
            add_bullet_list(doc, safe_names)
        else:
            add_paragraph(doc, "No annexes uploaded.")

        # -----------------
        # Save file to disk
        # -----------------
        # Build filename: include division and timestamp to avoid collisions
        div_part = (cover.get("Division") or "division").strip().replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"workplan_{div_part}_{timestamp}.docx"
        filepath = os.path.join(LOCAL_DATA_DIR, filename)

        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        # Save the docx
        doc.save(filepath)

        # -----------------
        # Attempt to push to GitHub (best-effort)
        # -----------------
        push_msg = ""
        try:
            # Save to a folder on repo, e.g. workplans/
            gh_path = f"workplans/{filename}"
            success, push_msg = push_file_to_github(filepath, gh_path)
            if not success:
                push_msg = f"GitHub push failed: {push_msg}"
            else:
                push_msg = f"Pushed to GitHub: {gh_path}"
        except Exception as e:
            push_msg = f"Pushing to GitHub failed: {e}"

        # Return values expected by finish_and_save()
        return (filepath, filename, push_msg)

    except Exception as e:
        # If anything fails, return falsy filepath and useful error message
        err = f"export_word failed: {e}\n{traceback.format_exc()}"
        return (None, None, err)


# ------------------------------------------------
# Save annexes immediately and persist their saved paths
# ------------------------------------------------
import hashlib

def _hash_bytes(b: bytes) -> str:
    h = hashlib.md5()
    h.update(b)
    return h.hexdigest()

# Save annexes without placeholders or duplicates
# --------------------------------------------------
def save_annexes_immediate(uploaded_files):
    if "annex_saved_list" not in st.session_state:
        st.session_state.annex_saved_list = []
    if "annex_saved_hashes" not in st.session_state:
        st.session_state.annex_saved_hashes = set()
    if "annex_saved_gh_paths" not in st.session_state:
        st.session_state.annex_saved_gh_paths = set()

    saved_results = []

    # Set of previously saved original names
    existing_names = {orig for (orig, *_ ) in st.session_state.annex_saved_list}

    for f in uploaded_files:
        try:
            # Skip if filename already saved
            if f.name in existing_names:
                saved_results.append((f.name, None, True, "Already saved — skipped"))
                continue

            # Load content
            b = f.getbuffer().tobytes() if hasattr(f, "getbuffer") else f.read()

            # Hash for duplicate detection
            content_hash = _hash_bytes(b)

            if content_hash in st.session_state.annex_saved_hashes:
                saved_results.append((f.name, None, True, "Already saved (identical content) — skipped"))
                continue

            # Build save path
            out_path = os.path.join(ANNEX_DIR, f.name)
            os.makedirs(os.path.dirname(out_path), exist_ok=True)

            # Write file ONCE
            with open(out_path, "wb") as out:
                out.write(b)

            # Push to GitHub only once per file
            gh_path = f"annexes/{f.name}"
            if USE_GITHUB and GITHUB_TOKEN and gh_path not in st.session_state.annex_saved_gh_paths:
                success, msg = push_file_to_github(out_path, gh_path)
                gh_msg = "Pushed to GitHub" if success else f"GitHub push failed: {msg}"
                if success:
                    st.session_state.annex_saved_gh_paths.add(gh_path)
            else:
                gh_msg = "Saved locally (GitHub skipped)."

            # Record FINAL entry only once
            st.session_state.annex_saved_list.append((f.name, out_path, content_hash, gh_path))
            st.session_state.annex_saved_hashes.add(content_hash)

            saved_results.append((f.name, out_path, True, gh_msg))

        except Exception as e:
            saved_results.append((f.name, None, False, f"Failed to save: {e}"))

    return saved_results


# ------------------------------------------------
# Finish callback: export docx + save master log + store filename for download
# ------------------------------------------------
def finish_and_save():

    # Prevent duplicate execution
    if st.session_state.get("finish_ran", False):
        return

    try:
        # Correct annex saving
        st.session_state.submission["Annexes_Saved"] = [
            item["path"] for item in st.session_state.annex_saved_list
        ]

        # Prevent repeated Word export
        if "word_generated" not in st.session_state:
            filepath, filename, push_result = export_word(st.session_state.submission)
            st.session_state.word_generated = True
        else:
            filepath = st.session_state.last_file
            push_result = st.session_state.last_push_result

        if not filepath:
            st.session_state.last_file = None
            st.session_state.finish_msg = f"Failed to generate Word doc: {push_result}"
            return

        st.session_state.generated_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        save_ok, save_msg = save_to_master_excel({
            "timestamp": datetime.now(),
            "division": st.session_state.submission.get("Cover", {}).get("Division", ""),
            "goals": ", ".join(st.session_state.submission.get("Selected Goals", [])),
            "data": str(st.session_state.submission)
        })

        st.session_state.last_file = filepath
        st.session_state.last_push_result = push_result

        if save_ok:
            st.session_state.finish_msg = save_msg or "Report generated successfully."
        else:
            st.session_state.finish_msg = save_msg or "Report generated, but master log save failed."

    except Exception as e:
        st.session_state.last_file = None
        st.session_state.finish_msg = f"Unexpected error: {e}\n{traceback.format_exc()}"

    # Mark finish as completed
    st.session_state.finish_ran = True


# ------------------------------------------------
# APP: Steps 1–9 (UI) — FULL PERSISTENCE (Option A)
# ------------------------------------------------

# ----------------------------
# STEP 1 — Division Cover Page
# ----------------------------
if st.session_state.step == 1:
    st.title("Step 1 — Division Workplan Cover Page")

    # Widgets with keys (widget state maintained by Streamlit)
    division = st.text_input("Division Name", key="division")
    director = st.text_input("Director's Name", key="director")
    date_entry = st.date_input("Date of Workplan", key="date_entry")
    version = st.text_input("Version of Workplan", key="version")
    ftes = st.text_input("Divisional FTEs", key="ftes")
    financial = st.text_input("Divisional Financial Resources", key="financial")
    signature = st.radio("Director's Signature Provided?", ["Yes", "No"], key="signature")

    # Persist ALWAYS so values reappear when navigating back
    st.session_state.submission["Cover"] = {
        "Division": division,
        "Director": director,
        "Date": str(date_entry),
        "Version": version,
        "FTEs": ftes,
        "Financial Resources": financial,
        "Director Signature": signature
    }

    # Navigation
    col1, col2 = st.columns([1, 1])
    with col1:
        st.write("")  # placeholder — no Previous on step 1
    with col2:
        st.button("Next", on_click=next_step, key="next_1")

# ----------------------------
# STEP 2 — Strategic Goals
# ----------------------------
if st.session_state.step == 2:
    st.title("Step 2 — Select Strategic Goals")

    try:
        df = pd.read_excel(ALIGNMENT_FILE)
    except Exception as e:
        st.error(f"Failed to read {ALIGNMENT_FILE}: {e}")
        st.stop()

    if "strategic_goal" not in df.columns:
        st.error(f"{ALIGNMENT_FILE} must contain a 'strategic_goal' column.")
        st.stop()

    goals = sorted(df["strategic_goal"].dropna().unique().tolist())

    # default from saved submission so selections reappear
    default_goals = st.session_state.submission.get("Selected Goals", [])

    selected_goals = st.multiselect(
        "Select Strategic Goals",
        options=goals,
        default=default_goals,
        key="selected_goals"
    )

    # persist immediately
    st.session_state.submission["Selected Goals"] = selected_goals

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step, key="prev_2")
    with col2:
        st.button("Next", on_click=next_step, key="next_2")

# ----------------------------
# STEP 3 — Aggregate Objectives + Other (dynamic)
# ----------------------------
if st.session_state.step == 3:
    st.title("Step 3 — Aggregate Divisional Objectives")
    df = pd.read_excel(ALIGNMENT_FILE)

    selected_goals = st.session_state.submission.get("Selected Goals", [])
    goal_to_agg = {}

    for g_idx, g in enumerate(selected_goals):
        st.subheader(f"Strategic Goal: {g}")

        agg_list = df[df["strategic_goal"] == g]["aggregate_divisional_objectives"].dropna().unique().tolist()

        # recover old values (both standard + custom)
        old_selected = st.session_state.submission.get("Aggregate Objectives", {}).get(g, [])
        default_standard = [x for x in old_selected if x in agg_list]

        sel = st.multiselect(
            f"Select Aggregate Objectives for {g}",
            options=agg_list,
            default=default_standard,
            key=f"agg_{g_idx}"
        )

        # Custom objectives — full persistence
        st.write("Custom Aggregate Objectives:")

        old_custom = [x for x in old_selected if x not in agg_list]
        prev_num_custom = len(old_custom)

        num_custom = st.number_input(
            f"How many custom aggregate objectives for {g}?",
            min_value=0,
            value=prev_num_custom,
            step=1,
            key=f"num_custom_{g_idx}"
        )

        custom_items = []
        for i in range(int(num_custom)):
            default_val = old_custom[i] if i < len(old_custom) else ""
            txt = st.text_input(
                f"Custom Objective {i+1} for {g}",
                value=default_val,
                key=f"custom_{g_idx}_{i}"
            )
            if txt.strip():
                custom_items.append(txt.strip())

        goal_to_agg[g] = sel + custom_items

    # save always
    st.session_state.submission["Aggregate Objectives"] = goal_to_agg

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step, key="prev_3")
    with col2:
        st.button("Next", on_click=next_step, key="next_3")

# ----------------------------
# STEP 4 — Activities & Results
# ----------------------------
if st.session_state.step == 4:
    st.title("Step 4 — Activities & Results")

    act_map = st.session_state.submission.get("Activities", {})
    new_map = {}

    aggregate_objectives = st.session_state.submission.get("Aggregate Objectives", {})
    if not aggregate_objectives:
        st.warning("No aggregate objectives found. Please go back to Step 3 and add them.")

    for g_idx, (g, agg_list) in enumerate(aggregate_objectives.items()):
        st.subheader(f"Strategic Goal: {g}")
        for a_idx, agg in enumerate(agg_list):
            st.markdown(f"### Aggregate Objective: {agg}")

            key_act = f"act_{g_idx}_{a_idx}"
            key_res = f"res_{g_idx}_{a_idx}"

            prev_vals = act_map.get((g, agg), {"activities": [], "results": []})
            activities_text = "\n".join(prev_vals.get("activities", []))
            results_text = "\n".join(prev_vals.get("results", []))

            activities = st.text_area(
                f"Planned activities (one per line) for '{agg}':",
                value=activities_text,
                key=key_act
            )
            results = st.text_area(
                f"Expected results (one per line) for '{agg}':",
                value=results_text,
                key=key_res
            )

            new_map[(g, agg)] = {
                "activities": [x.strip() for x in activities.split("\n") if x.strip()],
                "results": [x.strip() for x in results.split("\n") if x.strip()]
            }

    st.session_state.submission["Activities"] = new_map

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step, key="prev_5")
    with col2:
        st.button("Next", on_click=next_step, key="next_5")

# ----------------------------
# STEP 5 — Metrics per Strategic Goal
# ----------------------------
if st.session_state.step == 5:
    st.title("Step 5 — Metrics per Strategic Goal")

    old_metrics = st.session_state.submission.get("Goal Metrics", {})
    metrics = {}

    selected_goals = st.session_state.submission.get("Selected Goals", [])
    if not selected_goals:
        st.warning("No strategic goals selected. Please go back to Step 2 to select goals.")

    for g_idx, g in enumerate(selected_goals):
        st.subheader(f"Strategic Goal: {g}")
        old = old_metrics.get(g, {})

        fte = st.text_input(f"FTEs for {g}", value=old.get("FTEs", ""), key=f"fte_{g_idx}")
        fin = st.text_input(f"Financial Resources for {g}", value=old.get("Financial Resources", ""), key=f"fin_{g_idx}")
        kpi = st.text_area(f"Key Performance Indicators for {g}", value=old.get("KPIs", ""), key=f"kpi_{g_idx}")
        other = st.text_area(f"Other Metrics for {g}", value=old.get("Other Metrics", ""), key=f"other_{g_idx}")

        metrics[g] = {
            "FTEs": fte,
            "Financial Resources": fin,
            "KPIs": kpi,
            "Other Metrics": other
        }

    st.session_state.submission["Goal Metrics"] = metrics

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step, key="prev_6")
    with col2:
        st.button("Next", on_click=next_step, key="next_6")

# ----------------------------
# STEP 6 — Optional Objective/Result Metrics
# ----------------------------
if st.session_state.step == 6:
    st.title("Step 6 — Objective & Result Metrics (Optional)")

    opt = st.radio("Would you like to report metrics for objectives/results?",
                   ["No", "Yes"], key="opt_obj_res")

    obj_res_metrics = {}

    if opt == "Yes":
        activities_map = st.session_state.submission.get("Activities", {})

        for (g, agg), data in activities_map.items():
            st.subheader(f"Aggregate Objective: {agg}")

            # Metrics for the aggregate objective
            fte_agg = st.text_input(f"FTEs — Aggregate Objective '{agg}'", key=f"fte_agg_{g}_{agg}")
            fin_agg = st.text_input(f"Financial Resources — Aggregate Objective '{agg}'", key=f"fin_agg_{g}_{agg}")
            kpi_agg = st.text_area(f"KPIs — Aggregate Objective '{agg}'", key=f"kpi_agg_{g}_{agg}")
            other_agg = st.text_area(f"Other Metrics — Aggregate Objective '{agg}'", key=f"other_agg_{g}_{agg}")

            obj_res_metrics[(g, agg, "AGGREGATE")] = {
                "FTEs": fte_agg,
                "Financial Resources": fin_agg,
                "KPIs": kpi_agg,
                "Other Metrics": other_agg,
            }

            # Metrics per expected result
            for res in data.get("results", []):
                st.markdown(f"### Expected Result: {res}")

                fte = st.text_input(f"FTEs for '{res}'", key=f"fte_res_{g}_{agg}_{res}")
                fin = st.text_input(f"Financial Resources for '{res}'", key=f"fin_res_{g}_{agg}_{res}")
                kpi = st.text_area(f"KPIs for '{res}'", key=f"kpi_res_{g}_{agg}_{res}")
                other = st.text_area(f"Other Metrics for '{res}'", key=f"other_res_{g}_{agg}_{res}")

                obj_res_metrics[(g, agg, f"RES_{res}")] = {
                    "FTEs": fte,
                    "Financial Resources": fin,
                    "KPIs": kpi,
                    "Other Metrics": other,
                }

    st.session_state.submission["Objective/Result Metrics"] = obj_res_metrics

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step)
    with col2:
        st.button("Next", on_click=next_step)


# ----------------------------
# STEP 7 — Additional Information
# ----------------------------
if st.session_state.step == 7:
    st.title("Step 7 — Additional Information")

    old = st.session_state.submission.get("Additional", {})

    additional_info = {
        "Partnerships": st.text_area("Partnerships", value=old.get("Partnerships", ""), key="add_partnerships"),
        "Events": st.text_area("Events", value=old.get("Events", ""), key="add_events"),
        "Knowledge Products": st.text_area("Knowledge Products", value=old.get("Knowledge Products", ""), key="add_products"),
        "Knowledge Management": st.text_area("Knowledge Management Practices", value=old.get("Knowledge Management", ""), key="add_km"),
        "Cross-Divisional Initiatives": st.text_area("Participation in cross-divisional initiatives", value=old.get("Cross-Divisional Initiatives", ""), key="add_cross"),
        "Projects/Networks": st.text_area("Projects or Networks", value=old.get("Projects/Networks", ""), key="add_projects"),
        "Risks": st.text_area("Risks", value=old.get("Risks", ""), key="add_risks"),
        "Other Information": st.text_area("Other Information", value=old.get("Other Information", ""), key="add_other")
    }

    st.session_state.submission["Additional"] = additional_info

    col1, col2 = st.columns(2)
    with col1:
        st.button("Previous", on_click=prev_step, key="prev_8")
    with col2:
        st.button("Next", on_click=next_step, key="next_8")

# ----------------------------
# STEP 8 — Upload Annexes & Export
# ----------------------------
if st.session_state.step == 8:
    st.title("Step 8 — Upload Annexes & Export")

    # Initialize annex tracking like app_local.py
    if "annexes_saved" not in st.session_state:
        st.session_state.annexes_saved = False
    if "annex_saved_list" not in st.session_state:
        st.session_state.annex_saved_list = []

    uploaded_files = st.file_uploader(
        "Upload annex files (PDF, Word, Excel, images, etc.)",
        accept_multiple_files=True,
        key="annex_uploads"
    )

    os.makedirs(ANNEX_DIR, exist_ok=True)
    saved_files = []

    # -------------------------------------------------------------
    # SAVE ONLY ONCE — identical logic to app_local.py
    # -------------------------------------------------------------
    if uploaded_files and not st.session_state.annexes_saved:

        for file in uploaded_files:

            # deterministic saved filename: timestamp + original filename
            new_name = datetime.now().strftime("%Y%m%d_%H%M%S_") + file.name
            save_path = os.path.join(ANNEX_DIR, new_name)

            with open(save_path, "wb") as f:
                f.write(file.getbuffer())

            # Save metadata for later use in export
            saved_files.append({
                "original_name": file.name,
                "saved_name": new_name,
                "path": save_path
            })

            # PUSH ONLY ONCE TO GITHUB
            if USE_GITHUB and GITHUB_TOKEN:
                gh_path = f"annexes/{new_name}"
                push_file_to_github(save_path, gh_path)

        st.session_state.annex_saved_list = saved_files
        st.session_state.submission["Annexes_Saved"] = saved_files
        st.session_state.annexes_saved = True  # <--- prevents duplicates

        st.success(f"Saved {len(saved_files)} annex(es).")

    elif st.session_state.annexes_saved:
        st.info("Annexes already saved. Upload again to replace them.")

    # -------------------------------------------------------
    # Show list of annexes already saved
    # -------------------------------------------------------
    if st.session_state.annex_saved_list:
        st.subheader("Attached Annexes:")
        for a in st.session_state.annex_saved_list:
            st.write(f"• {a['original_name']}")

    st.write("---")

    # Navigation & Finish
    col1, col2 = st.columns(2)

    with col1:
        st.button("Previous", on_click=prev_step, key="prev_8")

    with col2:
        st.button("Finish & Generate Report", on_click=finish_and_save, key="finish_8")

    # -------------------------------------------------------
    # Download generated report
    # -------------------------------------------------------
    if st.session_state.last_file:
        st.success("✔ Workplan generated successfully!")

        try:
            with open(st.session_state.last_file, "rb") as f:
                st.download_button(
                    label="📥 Download Word Document",
                    data=f,
                    file_name=os.path.basename(st.session_state.last_file),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_btn_step8"
                )
        except Exception as e:
            st.error(f"File generated but download failed: {e}")
